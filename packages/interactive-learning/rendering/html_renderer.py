#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
interactive_learning/rendering/html_renderer.py
================================================

إعادة إنتاج الصفحة بصيغة HTML مع الحفاظ على التخطيط الكامل.
"""

import base64
import json
from dataclasses import asdict
from pathlib import Path
from typing import Dict, List, Optional, Union
import uuid

import cv2
import numpy as np
from jinja2 import Template


class HTMLRenderer:
    """
    محول للتخطيط إلى HTML تفاعلي.
    
    يحافظ على:
    - الترتيب المكاني الدقيق
    - الجداول مع تنسيقها
    - الرسومات كـ SVG
    - النصوص القابلة للتحرير
    """
    
    def __init__(
        self,
        include_interactive: bool = True,
        rtl: bool = True
    ):
        self.interactive = include_interactive
        self.rtl = rtl
        self.scale = 1.0  # عامل التكبير
    
    def render(
        self,
        layout,
        original_image: Optional[np.ndarray] = None,
        corrections: Optional[Dict[str, str]] = None,
        output_path: Optional[Path] = None
    ) -> str:
        """
        إنتاج HTML من التخطيط.
        
        Args:
            layout: PageLayout
            original_image: الصورة الأصلية (للخلفية)
            corrections: تصحيحات المستخدم {word_id: text}
            output_path: مسار الحفظ
        
        Returns:
            نص HTML
        """
        corrections = corrections or {}
        
        # تحويل الصورة لـ base64 إذا موجودة
        bg_image = None
        if original_image is not None:
            bg_image = self._image_to_base64(original_image)
        
        # بناء العناصر
        elements_html = []
        
        # الفقرات
        for para in layout.paragraphs:
            para_html = self._render_paragraph(para, corrections)
            elements_html.append(para_html)
        
        # الجداول
        for table in layout.tables:
            table_html = self._render_table(table, corrections)
            elements_html.append(table_html)
        
        # الرسومات
        for graphic in layout.graphics:
            graphic_html = self._render_graphic(graphic)
            elements_html.append(graphic_html)
        
        # تجميع الصفحة
        html = self._build_page(
            layout=layout,
            elements=elements_html,
            bg_image=bg_image
        )
        
        # الحفظ
        if output_path:
            output_path.write_text(html, encoding='utf-8')
        
        return html
    
    def _render_paragraph(
        self,
        paragraph,
        corrections: Dict[str, str]
    ) -> str:
        """تصيير فقرة."""
        lines_html = []
        
        for line in paragraph.lines:
            words_html = []
            
            for word in line.words:
                # تطبيق التصحيحات
                display_text = corrections.get(word.id, word.text)
                is_corrected = word.id in corrections
                
                # تحديد الأنماط
                classes = ['word']
                if is_corrected:
                    classes.append('corrected')
                if word.confidence < 0.7:
                    classes.append('low-confidence')
                
                # موقع الكلمة
                x1, y1, x2, y2 = word.bbox
                
                word_html = f'''
                <span class="{' '.join(classes)}"
                      data-word-id="{word.id}"
                      data-original="{word.text}"
                      data-confidence="{word.confidence:.3f}"
                      style="
                          position: absolute;
                          left: {x1}px;
                          top: {y1}px;
                          width: {x2 - x1}px;
                          height: {y2 - y1}px;
                          font-size: {self._estimate_font_size(y2 - y1)}px;
                          {'color: #2e7d32; font-weight: bold;' if is_corrected else ''}
                          {'border-bottom: 2px wavy #c62828;' if word.confidence < 0.7 else ''}
                      "
                      {'contenteditable="true"' if self.interactive else ''}
                      title="الأصلي: {word.text} | الثقة: {word.confidence:.1%}"
                >{display_text}</span>
                '''
                words_html.append(word_html)
            
            # تجميع السطر
            line_html = f'''
            <div class="line" style="
                position: absolute;
                left: {line.bbox[0]}px;
                top: {line.bbox[1]}px;
                width: {line.bbox[2] - line.bbox[0]}px;
                height: {line.bbox[3] - line.bbox[1]}px;
                text-align: {line.alignment};
                direction: {line.direction};
            ">
                {''.join(words_html)}
            </div>
            '''
            lines_html.append(line_html)
        
        # تجميع الفقرة
        return f'''
        <div class="paragraph" style="
            position: absolute;
            left: {paragraph.bbox[0]}px;
            top: {paragraph.bbox[1]}px;
            width: {paragraph.bbox[2] - paragraph.bbox[0]}px;
            min-height: {paragraph.bbox[3] - paragraph.bbox[1]}px;
        ">
            {''.join(lines_html)}
        </div>
        '''
    
    def _render_table(
        self,
        table,
        corrections: Dict[str, str]
    ) -> str:
        """تصيير جدول."""
        # بناء شبكة الخلايا
        rows = {}
        for cell in table.cells:
            if cell.row not in rows:
                rows[cell.row] = {}
            rows[cell.row][cell.col] = cell
        
        # بناء HTML
        rows_html = []
        for row_idx in sorted(rows.keys()):
            cells_html = []
            
            for col_idx in sorted(rows[row_idx].keys()):
                cell = rows[row_idx][col_idx]
                
                # تجميع نص الخلية
                cell_text = ' '.join(
                    corrections.get(w.id, w.text)
                    for w in cell.content
                )
                
                # أنماط الخلية
                cell_style = f'''
                    width: {(cell.bbox[2] - cell.bbox[0]) / (table.bbox[2] - table.bbox[0]) * 100}%;
                    height: {cell.bbox[3] - cell.bbox[1]}px;
                '''
                
                if cell.is_header:
                    cell_style += 'background: #f5f5f5; font-weight: bold;'
                
                if cell.colspan > 1:
                    cell_style += f' grid-column: span {cell.colspan};'
                
                cells_html.append(f'''
                    <td style="{cell_style}"
                        {'class="header"' if cell.is_header else ''}>
                        {cell_text}
                    </td>
                ''')
            
            rows_html.append(f'<tr>{"".join(cells_html)}</tr>')
        
        return f'''
        <table class="detected-table" style="
            position: absolute;
            left: {table.bbox[0]}px;
            top: {table.bbox[1]}px;
            width: {table.bbox[2] - table.bbox[0]}px;
            border-collapse: collapse;
            border: 2px solid #333;
        ">
            {''.join(rows_html)}
        </table>
        '''
    
    def _render_graphic(self, graphic) -> str:
        """تصيير عنصر رسومي كـ SVG."""
        x1, y1, x2, y2 = graphic.bbox
        width = x2 - x1
        height = y2 - y1
        
        if graphic.element_type == "diagram_box":
            return self._render_diagram_box(graphic, width, height)
        elif graphic.element_type == "chart":
            return self._render_chart(graphic, width, height)
        elif graphic.element_type == "arrow":
            return self._render_arrow(graphic, width, height)
        else:
            # صورة عادية
            return self._render_image_placeholder(graphic, width, height)
    
    def _render_diagram_box(
        self,
        graphic,
        width: int,
        height: int
    ) -> str:
        """تصيير صندوق مخطط تدفق."""
        render_data = graphic.render_data
        shape = render_data.get('shape', 'process_box')
        fill_color = render_data.get('fill_color', '#e3f2fd')
        border_color = render_data.get('border_color', '#1565c0')
        
        # النص
        text = ' '.join(w.text for w in graphic.detected_text)
        
        # إنشاء SVG حسب الشكل
        if shape == 'decision_diamond':
            # معين للقرار
            points = f"{width/2},0 {width},{height/2} {width/2},{height} 0,{height/2}"
            svg_content = f'''
                <polygon points="{points}"
                         fill="{fill_color}"
                         stroke="{border_color}"
                         stroke-width="2"/>
            '''
        elif shape == 'start_oval':
            # بيضاوي للبداية/النهاية
            svg_content = f'''
                <ellipse cx="{width/2}" cy="{height/2}"
                         rx="{width/2 - 5}" ry="{height/2 - 5}"
                         fill="{fill_color}"
                         stroke="{border_color}"
                         stroke-width="2"/>
            '''
        else:
            # مستطيل للعملية
            svg_content = f'''
                <rect x="5" y="5"
                      width="{width - 10}" height="{height - 10}"
                      rx="5" ry="5"
                      fill="{fill_color}"
                      stroke="{border_color}"
                      stroke-width="2"/>
            '''
        
        # إضافة نص
        svg_content += f'''
            <text x="{width/2}" y="{height/2}"
                  text-anchor="middle"
                  dominant-baseline="middle"
                  font-family="Segoe UI, Arial"
                  font-size="{min(width, height) // 8}"
                  fill="#333">
                {text}
            </text>
        '''
        
        return f'''
        <div class="graphic diagram-box" style="
            position: absolute;
            left: {graphic.bbox[0]}px;
            top: {graphic.bbox[1]}px;
            width: {width}px;
            height: {height}px;
        ">
            <svg width="100%" height="100%" viewBox="0 0 {width} {height}">
                {svg_content}
            </svg>
        </div>
        '''
    
    def _render_chart(
        self,
        graphic,
        width: int,
        height: int
    ) -> str:
        """تصيير مخطط بياني."""
        # استخراج البيانات من النص
        # TODO: تحليل البيانات وإنشاء مخطط حقيقي
        
        # مخطط شريطي بسيط كمثال
        return f'''
        <div class="graphic chart" style="
            position: absolute;
            left: {graphic.bbox[0]}px;
            top: {graphic.bbox[1]}px;
            width: {width}px;
            height: {height}px;
            background: white;
            border: 1px solid #ddd;
            border-radius: 8px;
            padding: 10px;
        ">
            <canvas id="chart_{graphic.id}" width="{width - 20}" height="{height - 20}"></canvas>
            <script>
                (function() {{
                    const ctx = document.getElementById('chart_{graphic.id}').getContext('2d');
                    // رسم مخطط بسيط
                    // TODO: استخدام مكتبة Chart.js
                }})();
            </script>
        </div>
        '''
    
    def _render_arrow(
        self,
        graphic,
        width: int,
        height: int
    ) -> str:
        """تصيير سهم."""
        # تحديد الاتجاه
        render_data = graphic.render_data
        
        return f'''
        <div class="graphic arrow" style="
            position: absolute;
            left: {graphic.bbox[0]}px;
            top: {graphic.bbox[1]}px;
            width: {width}px;
            height: {height}px;
        ">
            <svg width="100%" height="100%" viewBox="0 0 {width} {height}">
                <defs>
                    <marker id="arrowhead_{graphic.id}" markerWidth="10" markerHeight="7"
                            refX="9" refY="3.5" orient="auto">
                        <polygon points="0 0, 10 3.5, 0 7" fill="#666"/>
                    </marker>
                </defs>
                <line x1="0" y1="{height/2}" x2="{width-10}" y2="{height/2}"
                      stroke="#666" stroke-width="2"
                      marker-end="url(#arrowhead_{graphic.id})"/>
            </svg>
        </div>
        '''
    
    def _render_image_placeholder(
        self,
        graphic,
        width: int,
        height: int
    ) -> str:
        """تصيير placeholder لصورة."""
        return f'''
        <div class="graphic image-placeholder" style="
            position: absolute;
            left: {graphic.bbox[0]}px;
            top: {graphic.bbox[1]}px;
            width: {width}px;
            height: {height}px;
            background: #f5f5f5;
            border: 2px dashed #ccc;
            display: flex;
            align-items: center;
            justify-content: center;
            color: #999;
        ">
            [صورة]
        </div>
        '''
    
    def _build_page(
        self,
        layout,
        elements: List[str],
        bg_image: Optional[str]
    ) -> str:
        """بناء صفحة HTML كاملة."""
        
        # اتجاه النص
        dir_attr = 'rtl' if self.rtl else 'ltr'
        lang_attr = 'ar' if self.rtl else 'en'
        
        template = Template('''
<!DOCTYPE html>
<html lang="{{ lang }}" dir="{{ dir }}">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>OmniFile - المستند المستخرج</title>
    <style>
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }
        
        body {
            font-family: 'Segoe UI', 'Traditional Arabic', 'Arial', sans-serif;
            background: #f0f0f0;
            padding: 20px;
            {{ 'direction: rtl;' if rtl else '' }}
        }
        
        .page-container {
            position: relative;
            width: {{ layout.width }}px;
            height: {{ layout.height }}px;
            margin: 0 auto;
            background: white;
            box-shadow: 0 2px 8px rgba(0,0,0,0.15);
            overflow: hidden;
        }
        
        .page-background {
            position: absolute;
            top: 0;
            left: 0;
            width: 100%;
            height: 100%;
            opacity: 0.3;
            pointer-events: none;
        }
        
        .content-layer {
            position: relative;
            width: 100%;
            height: 100%;
        }
        
        /* الكلمات */
        .word {
            display: inline-block;
            white-space: nowrap;
            cursor: pointer;
            transition: all 0.2s;
            line-height: 1.2;
        }
        
        .word:hover {
            background: rgba(25, 118, 210, 0.1);
            outline: 1px solid #1976d2;
        }
        
        .word.corrected {
            background: rgba(46, 125, 50, 0.1);
        }
        
        .word.low-confidence {
            border-bottom: 2px wavy #c62828;
        }
        
        .word:focus {
            outline: 2px solid #1976d2;
            background: rgba(25, 118, 210, 0.2);
        }
        
        /* الجداول */
        .detected-table {
            background: white;
        }
        
        .detected-table td {
            border: 1px solid #ccc;
            padding: 8px;
            text-align: {{ 'right' if rtl else 'left' }};
        }
        
        .detected-table .header {
            background: #f5f5f5;
            font-weight: bold;
        }
        
        /* الرسومات */
        .graphic {
            pointer-events: none;
        }
        
        .graphic.diagram-box {
            pointer-events: auto;
        }
        
        /* شريط الأدوات */
        .toolbar {
            position: fixed;
            bottom: 20px;
            left: 50%;
            transform: translateX(-50%);
            background: white;
            padding: 10px 20px;
            border-radius: 30px;
            box-shadow: 0 4px 12px rgba(0,0,0,0.15);
            display: flex;
            gap: 10px;
            z-index: 1000;
        }
        
        .toolbar button {
            padding: 8px 16px;
            border: none;
            border-radius: 20px;
            background: #1976d2;
            color: white;
            cursor: pointer;
            font-size: 14px;
            transition: all 0.2s;
        }
        
        .toolbar button:hover {
            background: #1565c0;
        }
        
        /* معلومات الكلمة */
        .word-info {
            position: fixed;
            top: 20px;
            right: 20px;
            background: white;
            padding: 15px;
            border-radius: 8px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.15);
            max-width: 300px;
            display: none;
        }
        
        .word-info.visible {
            display: block;
        }
        
        /* طباعة */
        @media print {
            body {
                background: white;
                padding: 0;
            }
            
            .page-container {
                box-shadow: none;
                margin: 0;
            }
            
            .toolbar {
                display: none;
            }
        }
    </style>
</head>
<body>
    <div class="page-container" id="page">
        {% if bg_image %}
        <img src="{{ bg_image }}" class="page-background" alt=""/>
        {% endif %}
        
        <div class="content-layer">
            {{ elements | join('\\n') | safe }}
        </div>
    </div>
    
    {% if interactive %}
    <!-- شريط الأدوات -->
    <div class="toolbar">
        <button onclick="saveCorrections()">💾 حفظ التصحيحات</button>
        <button onclick="exportDocument()">📄 تصدير</button>
        <button onclick="toggleBackground()">👁️ خلفية</button>
        <button onclick="window.print()">🖨️ طباعة</button>
    </div>
    
    <!-- معلومات الكلمة -->
    <div class="word-info" id="wordInfo">
        <h4>معلومات الكلمة</h4>
        <p>الأصلي: <span id="infoOriginal"></span></p>
        <p>الثقة: <span id="infoConfidence"></span></p>
        <p>المعرف: <span id="infoId"></span></p>
    </div>
    
    <script>
        // تتبع التصحيحات
        let corrections = {};
        
        // تفاعل الكلمات
        document.querySelectorAll('.word').forEach(word => {
            word.addEventListener('focus', function() {
                showWordInfo(this);
            });
            
            word.addEventListener('blur', function() {
                const newText = this.textContent.trim();
                const original = this.dataset.original;
                const wordId = this.dataset.wordId;
                
                if (newText !== original) {
                    corrections[wordId] = newText;
                    this.classList.add('corrected');
                    console.log('تصحيح:', wordId, original, '->', newText);
                }
            });
        });
        
        function showWordInfo(element) {
            document.getElementById('infoOriginal').textContent = element.dataset.original;
            document.getElementById('infoConfidence').textContent = 
                (parseFloat(element.dataset.confidence) * 100).toFixed(1) + '%';
            document.getElementById('infoId').textContent = element.dataset.wordId;
            document.getElementById('wordInfo').classList.add('visible');
        }
        
        function saveCorrections() {
            fetch('/api/save-corrections', {
                method: 'POST',
                headers: {'Content-Type': 'application/json'},
                body: JSON.stringify(corrections)
            })
            .then(r => r.json())
            .then(data => {
                alert('تم حفظ ' + Object.keys(corrections).length + ' تصحيح');
            });
        }
        
        function exportDocument() {
            const format = prompt('اختر الصيغة: html, docx, pdf', 'html');
            if (format) {
                window.location.href = '/api/export?format=' + format;
            }
        }
        
        let bgVisible = true;
        function toggleBackground() {
            bgVisible = !bgVisible;
            document.querySelector('.page-background').style.opacity = 
                bgVisible ? '0.3' : '0';
        }
        
        // اختصارات لوحة المفاتيح
        document.addEventListener('keydown', function(e) {
            if (e.ctrlKey && e.key === 's') {
                e.preventDefault();
                saveCorrections();
            }
        });
    </script>
    {% endif %}
</body>
</html>
        ''')
        
        return template.render(
            lang=lang_attr,
            dir=dir_attr,
            rtl=self.rtl,
            layout=layout,
            elements=elements,
            bg_image=bg_image,
            interactive=self.interactive
        )
    
    def _image_to_base64(self, image: np.ndarray) -> str:
        """تحويل صورة لـ base64."""
        _, buffer = cv2.imencode('.jpg', image)
        return 'data:image/jpeg;base64,' + base64.b64encode(buffer).decode()
    
    def _estimate_font_size(self, pixel_height: int) -> float:
        """تقدير حجم الخط من الارتفاع بالبكسل."""
        # تقريب: ارتفاع الخط ≈ 1.2 × حجم الخط بالبكسل
        return max(8, pixel_height / 1.2)
    
    def render_with_layout_preservation(
        self,
        layout,
        output_path: Path,
        format: str = 'html'
    ) -> Path:
        """
        تصيير مع الحفاظ الكامل على التخطيط.
        
        يدعم:
        - html: صفحة تفاعلية
        - docx: مستند Word
        - pdf: PDF عالي الجودة
        """
        if format == 'html':
            html = self.render(layout, output_path=output_path)
            return output_path
        
        elif format == 'docx':
            return self._render_to_docx(layout, output_path)
        
        elif format == 'pdf':
            return self._render_to_pdf(layout, output_path)
        
        else:
            raise ValueError(f"Unsupported format: {format}")
    
    def _render_to_docx(self, layout, output_path: Path) -> Path:
        """تصيير لمستند Word."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # إعداد الصفحة
            section = doc.sections[0]
            section.page_width = Inches(layout.width / 96)  # افتراض 96 DPI
            section.page_height = Inches(layout.height / 96)
            
            # إضافة الفقرات
            for para in layout.paragraphs:
                doc_para = doc.add_paragraph()
                
                # المحاذاة
                if para.lines and para.lines[0].alignment == 'right':
                    doc_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # الكلمات
                for line in para.lines:
                    for word in line.words:
                        run = doc_para.add_run(word.text + ' ')
                        run.font.size = Pt(self._estimate_font_size(
                            word.bbox[3] - word.bbox[1]
                        ))
                        
                        # تنسيقات
                        if word.is_bold:
                            run.bold = True
                        if word.is_italic:
                            run.italic = True
                        if word.is_underlined:
                            run.underline = True
            
            # حفظ
            doc.save(output_path)
            return output_path
            
        except ImportError:
            raise ImportError("Install python-docx: pip install python-docx")
    
    def _render_to_pdf(self, layout, output_path: Path) -> Path:
        """تصيير لـ PDF."""
        # استخدام WeasyPrint أو ReportLab
        # أولاً ننشئ HTML ثم نحوله
        html_path = output_path.with_suffix('.html')
        self.render(layout, output_path=html_path)
        
        try:
            from weasyprint import HTML, CSS
            
            html = HTML(filename=str(html_path))
            html.write_pdf(str(output_path))
            
            # حذف الملف المؤقت
            html_path.unlink()
            
            return output_path
            
        except ImportError:
            # fallback: استخدام wkhtmltopdf
            import subprocess
            subprocess.run([
                'wkhtmltopdf',
                '--page-width', str(layout.width / 96),
                '--page-height', str(layout.height / 96),
                str(html_path),
                str(output_path)
            ], check=True)
            
            html_path.unlink()
            return output_path
