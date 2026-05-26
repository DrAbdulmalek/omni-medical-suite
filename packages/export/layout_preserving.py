"""
layout_preserving.py — تصدير DOCX/HTML مع الحفاظ على التخطيط البصري.
يقبل layout_data بصيغة JSON ويُعيد مستنداً يحاكي الكتابة الحاسوبية.

يدعم تنسيقين:
1. التنسيق البسيط (layout_data): {"image_path": "...", "blocks": [...]}
2. الهيكل القياسي (normalized): {"metadata": {...}, "pages": [{"blocks": [...]}]}

المؤلف: Dr Abdulmalek Tamer Al-husseini
الترخيص: MIT
"""
import json
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def _set_rtl(paragraph):
    """ضبط اتجاه الفقرة RTL."""
    pPr = paragraph._element.get_or_add_pPr()
    pPr.set(qn('w:bidi'), '1')


def export_to_docx(layout_data: dict, output_path: str) -> str:
    """
    تصدير layout_data إلى ملف DOCX يحافظ على البنية البصرية.
    يُرجع مسار الملف المُنشأ.

    التنسيق البسيط:
        layout_data = {
            "image_path": "...",
            "blocks": [{"type": "paragraph", "bbox": [...], "text": "..."}, ...]
        }
    """
    doc = Document()

    # هوامش موحدة
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # نمط افتراضي RTL
    style = doc.styles['Normal']
    style.font.size = Pt(12)
    rPr = style.element.get_or_add_rPr()
    rPr.set(qn('w:rtl'), '1')

    for block in layout_data.get('blocks', []):
        btype = block.get('type', 'paragraph')

        if btype in ('paragraph', 'caption'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_rtl(p)
            run = p.add_run(block.get('text', ''))
            run.font.size = Pt(10 if btype == 'caption' else 12)
            if btype == 'caption':
                run.italic = True

        elif btype == 'header':
            p = doc.add_heading(block.get('text', ''), level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_rtl(p)

        elif btype == 'table':
            cells = block.get('cells', [])
            if not cells:
                continue
            rows, cols = len(cells), max(len(r) for r in cells)
            tbl = doc.add_table(rows=rows, cols=cols, style='Table Grid')
            tbl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for i, row in enumerate(cells):
                for j, cell_text in enumerate(row):
                    if j < cols:
                        c = tbl.cell(i, j)
                        c.text = ''
                        p = c.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        _set_rtl(p)
                        p.add_run(str(cell_text)).font.size = Pt(11)
            doc.add_paragraph()  # مسافة بعد الجدول

        elif btype == 'image':
            img_file = block.get('image_file', '')
            if img_file and os.path.exists(img_file):
                doc.add_picture(img_file, width=Inches(4.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_path)
    return output_path


def layout_to_docx(layout_json_path: str, output_docx: str) -> str:
    """
    تصدير النتائج إلى DOCX مع الحفاظ على التخطيط.
    يعمل على ملف JSON بالهيكل القياسي (metadata + pages + blocks).

    هذا هو التنسيق المُنتج من modules.vision.normalize.normalize_ocr_output().

    Args:
        layout_json_path: مسار ملف JSON بالهيكل القياسي
        output_docx: مسار ملف DOCX المطلوب

    Returns:
        مسار ملف DOCX المُنشأ
    """
    with open(layout_json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    doc = Document()

    # هوامش موحدة
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # نمط افتراضي RTL
    style = doc.styles['Normal']
    style.font.size = Pt(12)
    rPr = style.element.get_or_add_rPr()
    rPr.set(qn('w:rtl'), '1')

    for page in data.get("pages", []):
        page_w = page.get("width", 2480)
        page_h = page.get("height", 3508)

        for block in page.get("blocks", []):
            b_type = block.get("type", "paragraph")

            if b_type == "paragraph":
                rtl = block.get("direction", "").lower() == "rtl"
                p = doc.add_paragraph()
                if rtl:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _set_rtl(p)
                run = p.add_run(block.get("text", ""))
                run.font.size = Pt(12)

            elif b_type == "header":
                p = doc.add_heading(block.get("text", ''), level=2)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _set_rtl(p)

            elif b_type == "table":
                cells_struct = block.get("structure", {}).get("cells", [])
                if not cells_struct:
                    # محاولة استخدام cells البسيط (للتوافق)
                    simple_cells = block.get("cells", [])
                    if simple_cells:
                        rows, cols = len(simple_cells), max(len(r) for r in simple_cells)
                        tbl = doc.add_table(rows=rows, cols=cols, style='Table Grid')
                        tbl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        for i, row in enumerate(simple_cells):
                            for j, cell_text in enumerate(row):
                                if j < cols:
                                    c = tbl.cell(i, j)
                                    c.text = ''
                                    p = c.paragraphs[0]
                                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                    _set_rtl(p)
                                    p.add_run(str(cell_text)).font.size = Pt(11)
                        doc.add_paragraph()
                    continue

                # إعادة بناء صفوف/أعمدة من الهيكل القياسي
                rows_dict: dict[int, dict[int, str]] = {}
                for cell in cells_struct:
                    r = cell["row"]
                    c = cell["col"]
                    rows_dict.setdefault(r, {})[c] = cell["text"]

                if not rows_dict:
                    continue

                max_col = max(max(row.keys()) for row in rows_dict.values()) + 1
                num_rows = len(rows_dict)

                tbl = doc.add_table(
                    rows=num_rows, cols=max_col, style='Table Grid'
                )
                tbl.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                for r in sorted(rows_dict.keys()):
                    for c in rows_dict[r]:
                        cell = tbl.cell(r, c)
                        cell.text = ''
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        _set_rtl(p)
                        run = p.add_run(str(rows_dict[r][c]))
                        run.font.size = Pt(11)

                doc.add_paragraph()

            elif b_type == "image":
                img_file = block.get("image_file", "")
                if img_file and os.path.exists(img_file):
                    doc.add_picture(img_file, width=Inches(4.5))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # التسمية
                if "caption" in block:
                    caption = block["caption"]
                    caption_text = (
                        caption["text"]
                        if isinstance(caption, dict)
                        else str(caption)
                    )
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(caption_text)
                    run.font.size = Pt(10)
                    run.italic = True

    doc.save(output_docx)
    return output_docx


def ocr_result_to_layout(ocr_json: dict, image_path: str = "") -> dict:
    """
    تحويل مخرجات OCR القياسية إلى تنسيق layout_data.
    """
    layout = {"image_path": image_path, "blocks": []}
    for block in ocr_json.get('blocks', []):
        nb = {
            "type": block.get('type', 'paragraph'),
            "bbox": block.get('bbox', [0, 0, 1, 1]),
            "text": block.get('text', ''),
        }
        if nb["type"] == 'table':
            nb["cells"] = block.get('cells', [])
        elif nb["type"] == 'image':
            nb["image_file"] = block.get('image_file', '')
        layout["blocks"].append(nb)
    return layout


# === Compatibility class for OmniFile_v500_Colab ===
class LayoutPreservingExporter:
    """واجهة متوافقة مع الـ notebook — تغلف الدوال المستقلة."""
    @staticmethod
    def export(layout_data: dict, output_path: str) -> str:
        return export_to_docx(layout_data, output_path)
