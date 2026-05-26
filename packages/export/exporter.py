"""
OmniFile AI Processor — Document Exporter
===========================================
Source: arabic-ocr-pro/core/exporter.py

Exports OCR results to multiple formats with RTL support:
- Plain text (UTF-8 with RTL BOM)
- JSON (structured with bounding boxes and confidence)
- DOCX (Microsoft Word with RTL paragraph support)
- HTML (preserves document layout with RTL styling)
- Searchable PDF (image + invisible text overlay)
"""

from __future__ import annotations

import io
import json
import logging
from datetime import datetime
from pathlib import Path
from typing import Optional, Union

logger = logging.getLogger(__name__)


class DocumentExporter:
    """Exports OCR documents to various output formats.

    Supports text, JSON, DOCX, HTML, and searchable PDF output.
    Each format preserves Arabic RTL text and document structure
    as appropriate.

    The exporter accepts a document object that follows a simple protocol:
    - ``document.pages`` — iterable of page objects
    - Each page has ``page_number``, ``width``, ``height``, and ``blocks``
    - Each block has ``block_type``, ``get_text()``, ``tokens``, ``bbox``,
      and optional ``table_data``
    - ``document.metadata`` has ``filename``, ``file_size``,
      ``page_count``, ``processing_time``, ``engine_used``

    Attributes:
        rtl: Whether to mark text as RTL in output formats that support it.
    """

    def __init__(self, rtl: bool = True) -> None:
        """Initialize the document exporter.

        Args:
            rtl: Whether to enable RTL text direction in supported formats.
        """
        self.rtl = rtl

    # ------------------------------------------------------------------
    # Main export dispatcher
    # ------------------------------------------------------------------

    def export(
        self,
        document,
        output_path: str | Path,
        format_name: str = "txt",
        images: Optional[list] = None,
    ) -> str:
        """Export a document to the specified format.

        Args:
            document: Processed OCR document (protocol-compliant object).
            output_path: Output file path.
            format_name: Output format
                         (``'txt'``, ``'json'``, ``'docx'``, ``'html'``,
                          ``'pdf'``).
            images: Optional list of page images (for searchable PDF).

        Returns:
            Absolute path to the exported file.

        Raises:
            ValueError: If the format is not supported.
            RuntimeError: If export fails.
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        exporters = {
            "txt": self.export_text,
            "json": self.export_json,
            "docx": self.export_docx,
            "html": self.export_html,
            "pdf": self.export_pdf,
        }

        format_name = format_name.lower().lstrip(".")
        if format_name not in exporters:
            raise ValueError(
                f"Unsupported export format: '{format_name}'. "
                f"Supported formats: {list(exporters.keys())}"
            )

        try:
            exporters[format_name](document, output_path, images)
            logger.info(
                f"Exported document to {output_path} (format: {format_name})"
            )
            return str(output_path.resolve())
        except Exception as exc:
            raise RuntimeError(
                f"Failed to export to {format_name}: {exc}"
            ) from exc

    # ------------------------------------------------------------------
    # Plain text
    # ------------------------------------------------------------------

    def export_text(
        self,
        document,
        output_path: str | Path,
        images: Optional[list] = None,
    ) -> None:
        """Export document as plain text (UTF-8).

        Preserves page structure and block ordering.

        Args:
            document: Processed OCR document.
            output_path: Output file path.
            images: Ignored for text export.
        """
        lines: list[str] = []

        for page in document.pages:
            lines.append(f"{'=' * 60}")
            lines.append(f"Page {page.page_number}")
            lines.append(f"{'=' * 60}")
            lines.append("")

            for block in page.blocks:
                text = block.get_text().strip()
                if not text:
                    continue

                block_type = self._get_block_type_value(block)

                if block_type == "HEADING":
                    lines.append("")
                    lines.append(text)
                    lines.append(
                        "-" * len(text) if len(text) < 60 else "-" * 60
                    )
                    lines.append("")
                elif (
                    block_type == "TABLE"
                    and hasattr(block, "table_data")
                    and block.table_data
                ):
                    lines.append("[TABLE]")
                    for row in block.table_data:
                        line = " | ".join(cell.strip() for cell in row)
                        lines.append(line)
                    lines.append("[/TABLE]")
                    lines.append("")
                else:
                    lines.append(text)
                    lines.append("")

        content = "\n".join(lines)

        with open(output_path, "w", encoding="utf-8") as f:
            # Write BOM for better RTL support in some editors
            f.write("\ufeff")
            f.write(content)

    # ------------------------------------------------------------------
    # JSON
    # ------------------------------------------------------------------

    def export_json(
        self,
        document,
        output_path: str | Path,
        images: Optional[list] = None,
    ) -> None:
        """Export document as structured JSON.

        Includes all tokens, bounding boxes, confidence scores,
        and metadata.

        Args:
            document: Processed OCR document.
            output_path: Output file path.
            images: Ignored for JSON export.
        """
        meta = document.metadata
        data = {
            "metadata": {
                "filename": getattr(meta, "filename", ""),
                "file_size": getattr(meta, "file_size", 0),
                "page_count": getattr(meta, "page_count", 0),
                "processing_time": getattr(meta, "processing_time", 0),
                "engine": getattr(meta, "engine_used", ""),
                "exported_at": datetime.now().isoformat(),
            },
            "pages": [],
        }

        for page in document.pages:
            page_data = {
                "page_number": page.page_number,
                "width": getattr(page, "width", 0),
                "height": getattr(page, "height", 0),
                "blocks": [],
            }

            for block in page.blocks:
                block_data = {
                    "type": self._get_block_type_value(block),
                    "text": block.get_text(),
                    "confidence": (
                        block.compute_confidence()
                        if hasattr(block, "compute_confidence")
                        else None
                    ),
                    "bbox": (
                        block.bbox.model_dump()
                        if hasattr(block, "bbox") and block.bbox
                        and hasattr(block.bbox, "model_dump")
                        else None
                    ),
                    "tokens": [
                        {
                            "text": token.text,
                            "confidence": token.confidence,
                            "bbox": (
                                token.bbox.model_dump()
                                if hasattr(token.bbox, "model_dump")
                                else None
                            ),
                            "engine": getattr(token, "engine", ""),
                        }
                        for token in block.tokens
                    ],
                }

                if (
                    self._get_block_type_value(block) == "TABLE"
                    and hasattr(block, "table_data")
                    and block.table_data
                ):
                    block_data["table_data"] = block.table_data

                page_data["blocks"].append(block_data)

            data["pages"].append(page_data)

        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

    # ------------------------------------------------------------------
    # DOCX
    # ------------------------------------------------------------------

    def export_docx(
        self,
        document,
        output_path: str | Path,
        images: Optional[list] = None,
    ) -> None:
        """Export document as a DOCX file with RTL support.

        Creates a Microsoft Word document with proper RTL paragraph
        formatting for Arabic text.

        Args:
            document: Processed OCR document.
            output_path: Output file path.
            images: Ignored for DOCX export.

        Raises:
            RuntimeError: If python-docx is not installed.
        """
        try:
            from docx import Document as DocxDocument
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise RuntimeError(
                "python-docx is required for DOCX export. "
                "Install with: pip install python-docx"
            )

        doc = DocxDocument()

        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Arial"
        font.size = Pt(11)

        for page in document.pages:
            for block in page.blocks:
                text = block.get_text().strip()
                if not text:
                    continue

                block_type = self._get_block_type_value(block)

                if block_type == "HEADING":
                    heading = doc.add_heading(level=2)
                    run = heading.add_run(text)
                    run.font.rtl = self.rtl
                    heading.paragraph_format.alignment = (
                        WD_ALIGN_PARAGRAPH.RIGHT
                        if self.rtl
                        else WD_ALIGN_PARAGRAPH.LEFT
                    )

                elif (
                    block_type == "TABLE"
                    and hasattr(block, "table_data")
                    and block.table_data
                ):
                    rows = len(block.table_data)
                    cols = (
                        max(len(row) for row in block.table_data)
                        if block.table_data
                        else 0
                    )
                    if rows > 0 and cols > 0:
                        table = doc.add_table(rows=rows, cols=cols)
                        table.style = "Table Grid"

                        for i, row_data in enumerate(block.table_data):
                            for j, cell_text in enumerate(row_data):
                                if j < cols:
                                    cell = table.cell(i, j)
                                    cell.text = cell_text.strip()

                        doc.add_paragraph("")

                else:
                    para = doc.add_paragraph()
                    run = para.add_run(text)
                    run.font.rtl = self.rtl
                    para.paragraph_format.alignment = (
                        WD_ALIGN_PARAGRAPH.RIGHT
                        if self.rtl
                        else WD_ALIGN_PARAGRAPH.LEFT
                    )

        doc.save(str(output_path))

    # ------------------------------------------------------------------
    # HTML
    # ------------------------------------------------------------------

    def export_html(
        self,
        document,
        output_path: str | Path,
        images: Optional[list] = None,
    ) -> None:
        """Export document as an HTML file preserving layout.

        Creates a styled HTML document with RTL support and
        document structure preserved.

        Args:
            document: Processed OCR document.
            output_path: Output file path.
            images: Optional page images (base64 encoded or file paths).
        """
        html_parts: list[str] = []

        html_parts.append("""<!DOCTYPE html>
<html lang="ar" dir="rtl">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>OCR Document</title>
    <style>
        body { font-family: 'Arial', 'Tahoma', sans-serif; direction: rtl; text-align: right; max-width: 800px; margin: 0 auto; padding: 20px; }
        .page { border: 1px solid #ccc; margin: 20px 0; padding: 20px; page-break-after: always; }
        .page-header { text-align: center; color: #666; border-bottom: 1px solid #eee; padding-bottom: 10px; margin-bottom: 20px; }
        .heading { font-size: 1.3em; font-weight: bold; margin: 15px 0; color: #333; }
        .paragraph { margin: 10px 0; line-height: 1.8; text-align: justify; }
        table { border-collapse: collapse; width: 100%; margin: 15px 0; }
        th, td { border: 1px solid #ddd; padding: 8px; text-align: right; }
        th { background-color: #f5f5f5; font-weight: bold; }
        tr:nth-child(even) { background-color: #f9f9f9; }
        .metadata { color: #999; font-size: 0.85em; margin-top: 30px; }
    </style>
</head>
<body>
""")

        for page in document.pages:
            html_parts.append('    <div class="page">')
            html_parts.append(
                f'        <div class="page-header">Page {page.page_number}</div>'
            )

            for block in page.blocks:
                text = block.get_text().strip()
                if not text:
                    continue

                # Escape HTML entities
                text = (
                    text.replace("&", "&amp;")
                    .replace("<", "&lt;")
                    .replace(">", "&gt;")
                    .replace('"', "&quot;")
                )

                block_type = self._get_block_type_value(block)

                if block_type == "HEADING":
                    html_parts.append(
                        f'        <div class="heading">{text}</div>'
                    )
                elif (
                    block_type == "TABLE"
                    and hasattr(block, "table_data")
                    and block.table_data
                ):
                    html_parts.append("        <table>")
                    for i, row in enumerate(block.table_data):
                        tag = "th" if i == 0 else "td"
                        cells = "".join(
                            f"<{tag}>{cell.strip()}</{tag}>"
                            for cell in row
                        )
                        html_parts.append(f"            <tr>{cells}</tr>")
                    html_parts.append("        </table>")
                else:
                    html_parts.append(
                        f'        <div class="paragraph">{text}</div>'
                    )

            html_parts.append("    </div>")

        # Add metadata footer
        meta = document.metadata
        html_parts.append(f"""
    <div class="metadata">
        <p>Generated by OmniFile AI Processor</p>
        <p>Source: {getattr(meta, 'filename', '')} | "
        f"Pages: {getattr(meta, 'page_count', '')} | "
        f"Engine: {getattr(meta, 'engine_used', '')}</p>
        <p>Processing time: {getattr(meta, 'processing_time', 0):.2f}s</p>
    </div>
</body>
</html>
""")

        content = "\n".join(html_parts)

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(content)

    # ------------------------------------------------------------------
    # Searchable PDF
    # ------------------------------------------------------------------

    def export_pdf(
        self,
        document,
        output_path: str | Path,
        images: Optional[list] = None,
    ) -> None:
        """Export document as a searchable PDF.

        If page images are provided, creates a PDF with the original
        images and an invisible text overlay for searchability.
        Otherwise, creates a text-only PDF.

        Args:
            document: Processed OCR document.
            output_path: Output file path.
            images: Optional list of page images (numpy arrays).

        Raises:
            RuntimeError: If fpdf2 is not installed.
        """
        try:
            from fpdf import FPDF
        except ImportError:
            raise RuntimeError(
                "fpdf2 is required for PDF export. "
                "Install with: pip install fpdf2"
            )

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)

        # Register a font that supports Arabic
        try:
            pdf.add_font(
                "DejaVu", "",
                "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                uni=True,
            )
            pdf.set_font("DejaVu", size=12)
        except RuntimeError:
            pdf.set_font("Helvetica", size=12)

        for page_idx, page in enumerate(document.pages):
            pdf.add_page()

            # If we have page images, add as background
            if images and page_idx < len(images):
                try:
                    img = images[page_idx]
                    import cv2

                    # Convert BGR to RGB
                    if len(img.shape) == 3:
                        img_rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
                    else:
                        img_rgb = cv2.cvtColor(img, cv2.COLOR_GRAY2RGB)

                    # Save as temporary PNG
                    temp_path = Path(output_path).with_suffix(
                        f".page{page_idx}.png"
                    )
                    cv2.imwrite(
                        str(temp_path),
                        cv2.cvtColor(img_rgb, cv2.COLOR_RGB2BGR),
                    )

                    # Calculate image placement
                    page_w = pdf.w - 20
                    img_h = (
                        (img_rgb.shape[0] / img_rgb.shape[1]) * page_w
                    )

                    if img_h > pdf.h - 20:
                        scale = (pdf.h - 20) / img_h
                        page_w = page_w * scale
                        img_h = pdf.h - 20

                    pdf.image(str(temp_path), x=10, y=10, w=page_w)

                    # Clean up temp file
                    temp_path.unlink(missing_ok=True)
                except Exception as exc:
                    logger.warning(
                        f"Failed to add page image to PDF: {exc}"
                    )

            # Add invisible text layer for searchability
            for block in page.blocks:
                text = block.get_text().strip()
                if not text:
                    continue

                # White text for invisible overlay (if image present)
                if images and page_idx < len(images):
                    pdf.set_text_color(255, 255, 255)
                else:
                    pdf.set_text_color(0, 0, 0)

                block_type = self._get_block_type_value(block)

                if block_type == "HEADING":
                    pdf.set_font_size(14)
                else:
                    pdf.set_font_size(11)

                # Estimate position from bbox
                if (
                    hasattr(block, "bbox")
                    and block.bbox
                    and images
                    and page_idx < len(images)
                ):
                    img = images[page_idx]
                    img_h_px, img_w_px = img.shape[:2]
                    x_pos = 10 + (block.bbox.x / img_w_px) * page_w
                    y_pos = 10 + (
                        block.bbox.y / img_h_px
                    ) * (pdf.h - 20)
                    pdf.set_xy(x_pos, y_pos)

                pdf.multi_cell(0, 6, text, align="R" if self.rtl else "L")

        pdf.output(str(output_path))

    # ------------------------------------------------------------------
    # Utilities
    # ------------------------------------------------------------------

    @staticmethod
    def get_supported_formats() -> list[str]:
        """Get list of supported export formats.

        Returns:
            List of format names (lowercase, without dots).
        """
        return ["txt", "json", "docx", "html", "pdf"]

    @staticmethod
    def detect_format(path: str | Path) -> str:
        """Detect the export format from a file path.

        Args:
            path: Output file path.

        Returns:
            Format name (lowercase, without dot).

        Raises:
            ValueError: If the format cannot be determined.
        """
        path = Path(path)
        ext = path.suffix.lower().lstrip(".")

        format_map = {
            "txt": "txt",
            "text": "txt",
            "json": "json",
            "docx": "docx",
            "doc": "docx",
            "html": "html",
            "htm": "html",
            "pdf": "pdf",
        }

        if ext in format_map:
            return format_map[ext]

        raise ValueError(
            f"Cannot determine export format from extension: '{ext}'"
        )

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _get_block_type_value(block) -> str:
        """Extract the string value of a block's type.

        Supports both enum-style (``block.block_type.value``) and
        plain-string ``block.block_type`` conventions.
        """
        bt = getattr(block, "block_type", None)
        if bt is None:
            return "TEXT"
        return getattr(bt, "value", str(bt))
