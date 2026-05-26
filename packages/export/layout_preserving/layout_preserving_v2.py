# layout_preserving.py - Layout-preserving export for OCR results

"""
وحدة التصدير مع الحفاظ على التخطيط الكامل - محاكاة الكتابة الحاسوبية.
يعتمد على ملف JSON يحتوي على بيانات التخطيط المستخرجة.
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image

def export_to_docx(layout_data, output_path):
    """
    إنشاء ملف DOCX من بيانات التخطيط.

    layout_data: dict يحوي:
        - image_path: مسار الصورة الأصلية (اختياري للرجوع إليها)
        - page_width, page_height (بالبكسل أو بالنسبة للصورة)
        - blocks: قائمة بالكتل، كل كتلة:
            {
                "type": "paragraph" | "table" | "image" | "caption",
                "bbox": [x1, y1, x2, y2] (إحداثيات نسبية 0-1),
                "text": "...",
                "cells": [[...]]  (للجداول فقط),
                "image_file": "img1.png" (للصور فقط)
            }
    """
    doc = Document()

    # إعداد الهوامش
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # افتراض أن جميع النصوص عربية RTL (يمكن تحسينه لاحقًا حسب اللغة)
    # إعداد النمط الافتراضي للـ RTL
    style = doc.styles['Normal']
    style.font.size = Pt(12)
    rPr = style.element.get_or_add_rPr()
    rPr.set(qn('w:rtl'), '1')

    # معالجة كل كتلة حسب النوع
    for block in layout_data.get('blocks', []):
        btype = block.get('type', 'paragraph')

        if btype == 'paragraph':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(block.get('text', ''))
            # ضبط الخط إن لزم
            run.font.size = Pt(12)

        elif btype == 'table':
            cells = block.get('cells', [])
            if not cells:
                continue
            rows = len(cells)
            cols = len(cells[0]) if rows > 0 else 0
            table = doc.add_table(rows=rows, cols=cols, style='Table Grid')
            table.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            for i, row_cells in enumerate(cells):
                for j, cell_text in enumerate(row_cells):
                    cell = table.cell(i, j)
                    # مسح المحتوى الافتراضي
                    cell.text = ''
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    run = p.add_run(str(cell_text))
                    run.font.size = Pt(11)
            # إضافة فراغ بعد الجدول
            doc.add_paragraph()

        elif btype == 'image':
            img_file = block.get('image_file')
            if img_file and os.path.exists(img_file):
                # يمكن تحديد العرض بناءً على bbox لكننا نستخدم حجم مناسب
                doc.add_picture(img_file, width=Inches(4.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif btype == 'caption':
            # تسمية الصورة تدرج مباشرة بعد الصورة (يجب أن يكون الترتيب صحيحًا في JSON)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(block.get('text', ''))
            run.font.size = Pt(10)
            run.italic = True
            # عادةً التسمية تكون RTL إذا كانت عربية
            pPr = p._element.get_or_add_pPr()
            pPr.set(qn('w:bidi'), '1')

    doc.save(output_path)
    print(f"✅ تم تصدير المستند إلى: {output_path}")


# === Compatibility class for OmniFile_v500_Colab ===
class LayoutPreservingExporter:
    """واجهة متوافقة مع الـ notebook — تغلف export_to_docx."""
    @staticmethod
    def export(layout_data: dict, output_path: str) -> str:
        return export_to_docx(layout_data, output_path)
